#!/usr/bin/env python3
"""Build a Word .docx from all guide chapters using Pandoc.

Usage:
    python make_docx.py

First run (no reference.docx present):
    Extracts Pandoc's default reference.docx template. Open it in Word,
    customize the styles (Heading 1–4, Normal, Code, Caption, etc.),
    save it back, then re-run to produce the final document.

Subsequent runs:
    Converts all chapters to a single .docx using the customized template.

After opening the output in Word:
    1. Accept the prompt to update fields — the TOC populates automatically.
    2. Inspect math equations — Pandoc converts LaTeX to OMML (Word equations).
       Complex expressions may need manual correction.
    3. Export to PDF: File > Export > Create PDF/XPS.

Requirements:
    pandoc    — https://pandoc.org/installing.html
    inkscape  — https://inkscape.org  (for SVG → PNG conversion)
"""

import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

ROOT = Path(__file__).parent
OUTPUT_DOCX = ROOT / "IFC_Alignment_Geometry_Implementation_Guide.docx"
REFERENCE_DOCX = ROOT / "reference.docx"

INKSCAPE = Path("C:/Program Files/Inkscape/bin/inkscape.exe")

CHAPTERS = [
    "Cover.md",
    "Foreword.md",
    "TableOfContents.md",
    "Notation.md",
    "1_IFC_Alignment_Concepts.md",
    "2_Horizontal_Alignments.md",
    "3_Vertical_Alignments.md",
    "4_Cant_Alignments.md",
    "5_OffsetCurves.md",
    "6_Approximate_Alignment_Geometry.md",
    "7_Alignments_Reusing_Horizontal_Layout.md",
    "8_LinearPlacement.md",
    "9_Referents_and_Stationing.md",
    "10_Sectioned_Surfaces_and_Solids.md",
    "11_Precision_and_Tolerance.md",
    "12_Alignment_Geometry_Testset.md",
    "Appendix_A_LandXML.md",
    "Index.md",
]

# Raw OpenXML TOC field injected into TableOfContents.md during docx preprocessing.
# Word updates it automatically on first open (w:dirty="true").
# TableOfContents.md itself is left unchanged so make_pdf.py still renders its static links.
# Front-matter and back-matter headings are assigned Title style in preprocessing so
# they are not Heading 1 and are therefore excluded from this field automatically.
_TOC_FIELD = (
    '\n\n```{=openxml}\n'
    '<w:p>'
    '<w:r><w:fldChar w:fldCharType="begin" w:dirty="true"/></w:r>'
    '<w:r><w:instrText xml:space="preserve"> TOC \\t "Heading 1,1" \\h \\z </w:instrText></w:r>'
    '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
    '<w:r><w:fldChar w:fldCharType="end"/></w:r>'
    '</w:p>\n'
    '```\n'
)



# Section break injected before Chapter 1.
# The w:sectPr ends the front matter section (Cover + Foreword + Notation)
# and sets its page number format to lowercase Roman (i, ii, iii…).
# The document-level sectPr (configured in add_page_numbers) governs
# Chapter 1 onwards with Arabic numbering restarting at 1.
_FRONT_MATTER_SECTION_BREAK = (
    '\n\n```{=openxml}\n'
    '<w:p><w:pPr><w:sectPr>'
    '<w:type w:val="nextPage"/>'
    '<w:pgNumType w:fmt="lowerRoman" w:start="0"/>'
    '</w:sectPr></w:pPr></w:p>\n'
    '```\n\n'
)


def check_pandoc() -> str:
    """Verify Pandoc is on PATH and return its version string."""
    try:
        result = subprocess.run(
            ["pandoc", "--version"],
            capture_output=True, text=True, check=True,
        )
        return result.stdout.splitlines()[0]
    except FileNotFoundError:
        print("ERROR: pandoc not found on PATH.")
        print("Install from: https://pandoc.org/installing.html")
        sys.exit(1)


def generate_reference_docx() -> None:
    """Extract Pandoc's built-in reference.docx as a customization starting point."""
    with open(REFERENCE_DOCX, "wb") as f:
        subprocess.run(
            ["pandoc", "--print-default-data-file", "reference.docx"],
            stdout=f, check=True,
        )
    print(f"Generated default reference template: {REFERENCE_DOCX}")
    print()
    print("Next steps:")
    print("  1. Open reference.docx in Word.")
    print("  2. Customize styles: Heading 1–4, Normal, Code Block, Caption,")
    print("     Verbatim Char (inline code), Block Text (blockquote).")
    print("  3. Set Heading 1 to 'Page break before' in paragraph settings.")
    print("  4. Save and close, then re-run this script.")


def convert_svgs(markdown_texts: list[str], png_dir: Path) -> list[str]:
    """Convert every SVG image reference to PNG, return updated markdown texts.

    Converted PNGs are written to png_dir. Inkscape is used for conversion.
    If Inkscape is unavailable, SVG references are left unchanged and a warning
    is printed (Pandoc will then skip them as it cannot embed SVGs directly).
    """
    if not INKSCAPE.exists():
        print(f"  WARNING: Inkscape not found at {INKSCAPE} — SVG images will be skipped.")
        return markdown_texts

    svg_pattern = re.compile(r'(!\[([^\]]*)\]\()([^)]+\.svg)(\))', re.IGNORECASE)
    converted: dict[str, str] = {}  # svg_path -> png_path

    def _convert(svg_rel: str) -> str:
        if svg_rel in converted:
            return converted[svg_rel]
        svg_abs = (ROOT / svg_rel).resolve()
        if not svg_abs.exists():
            print(f"    SKIP (not found): {svg_rel}")
            converted[svg_rel] = svg_rel
            return svg_rel
        png_name = svg_abs.stem + ".png"
        png_abs = png_dir / png_name
        if not png_abs.exists():
            result = subprocess.run(
                [str(INKSCAPE), f"--export-filename={png_abs}", str(svg_abs)],
                capture_output=True, text=True,
            )
            if result.returncode != 0 or not png_abs.exists():
                print(f"    WARN: Inkscape failed on {svg_rel}")
                converted[svg_rel] = svg_rel
                return svg_rel
        png_rel = str(png_abs)
        converted[svg_rel] = png_rel
        return png_rel

    updated = []
    for text in markdown_texts:
        def _repl(m: re.Match) -> str:
            return m.group(1) + _convert(m.group(3)) + m.group(4)
        updated.append(svg_pattern.sub(_repl, text))

    n = sum(1 for v in converted.values() if v != list(converted.keys())[list(converted.values()).index(v)] or v.endswith('.png'))
    print(f"  SVG->PNG: {sum(1 for k, v in converted.items() if v != k)} converted, "
          f"{sum(1 for k, v in converted.items() if v == k)} unchanged/skipped")
    return updated


def add_page_numbers(docx_path: Path) -> None:
    """Add page numbers to the footer:
    - Cover (first page of first section): suppressed.
    - Foreword through Notation: lowercase Roman numerals (i, ii, iii…).
    - Chapter 1 onwards: Arabic numerals restarting at 1.
    """
    if not _DOCX_AVAILABLE:
        print("  SKIP page numbers: python-docx not installed (pip install python-docx)")
        return

    doc = Document(str(docx_path))
    sections = list(doc.sections)

    for idx, section in enumerate(sections):
        # Suppress the page number on the first page of the first section (the cover).
        section.different_first_page_header_footer = (idx == 0)

        footer = section.footer
        for para in footer.paragraphs:
            para.clear()

        para = footer.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.add_run()
        for fld_type, instr in [("begin", None), (None, " PAGE "), ("end", None)]:
            if fld_type:
                el = OxmlElement("w:fldChar")
                el.set(qn("w:fldCharType"), fld_type)
            else:
                el = OxmlElement("w:instrText")
                el.set(qn("xml:space"), "preserve")
                el.text = instr
            run._r.append(el)

    # Set the last section (Chapter 1 onwards) to Arabic numbering restarting at 1.
    last_sectPr = sections[-1]._sectPr
    for existing in last_sectPr.findall(qn('w:pgNumType')):
        last_sectPr.remove(existing)
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), 'decimal')
    pgNumType.set(qn('w:start'), '1')
    last_sectPr.append(pgNumType)

    doc.save(str(docx_path))
    n_sections = len(sections)
    if n_sections >= 2:
        print(f"  Page numbers added ({n_sections} sections: Roman for front matter, Arabic from Chapter 1)")
    else:
        print("  Page numbers added (single section — section break may not have been recognised)")


def main() -> None:
    version = check_pandoc()
    print(f"Using {version}")

    missing = [f for f in CHAPTERS if not (ROOT / f).exists()]
    if missing:
        for f in missing:
            print(f"  WARNING: not found, skipping — {f}")

    chapter_paths = [ROOT / f for f in CHAPTERS if (ROOT / f).exists()]

    if not REFERENCE_DOCX.exists():
        print("No reference.docx found — generating default template...")
        generate_reference_docx()
        return

    tmp = Path(tempfile.mkdtemp(prefix="make_docx_"))
    try:
        print("Converting SVG images to PNG...")
        texts = [p.read_text(encoding="utf-8-sig") for p in chapter_paths]
        texts = convert_svgs(texts, tmp)

        page_break = '\n\n```{=openxml}\n<w:p><w:r><w:br w:type="page"/></w:r></w:p>\n```\n\n'

        # Front/back matter headings use non-Heading-1 styles so the TOC field
        # (which collects Heading 1 only) excludes them automatically.
        FRONTMATTER = {"Foreword.md", "Notation.md"}   # → Frontmatter Heading style
        BACKMATTER  = {"Index.md"}                      # → Title style

        # Replace TableOfContents.md static list with a live Word TOC field.
        # Apply appropriate styles to front/back matter headings.
        for i, path in enumerate(chapter_paths):
            if path.name == "TableOfContents.md":
                m = re.match(r'^# (.+)$', texts[i], re.MULTILINE)
                heading_text = m.group(1) if m else "Table of Contents"
                texts[i] = (
                    f':::{{custom-style="Frontmatter Heading"}}\n{heading_text}\n:::\n\n'
                    + _TOC_FIELD
                )
                print("  Injected live TOC field")
            elif path.name in FRONTMATTER:
                texts[i] = re.sub(
                    r'^# (.+)$',
                    lambda m: f':::{{custom-style="Frontmatter Heading"}}\n{m.group(1)}\n:::',
                    texts[i], flags=re.MULTILINE,
                )
            elif path.name in BACKMATTER:
                texts[i] = re.sub(
                    r'^# (.+)$',
                    lambda m: f':::{{custom-style="Title"}}\n{m.group(1)}\n:::',
                    texts[i], flags=re.MULTILINE,
                )
            elif path.name == "Cover.md":
                # Apply Word styles and image sizing to cover elements.
                # Done here rather than in the source markdown so Pandoc-specific
                # syntax doesn't appear as literal text in standard markdown renderers.
                t = texts[i]
                # # heading → Title style
                t = re.sub(
                    r'^# (.+)$',
                    lambda m: f':::{{custom-style="Title"}}\n{m.group(1)}\n:::',
                    t, flags=re.MULTILINE,
                )
                # First *...* paragraph → Subtitle style
                t = re.sub(
                    r'^\*([^*\n]+)\*[ \t]*$',
                    lambda m: f':::{{custom-style="Subtitle"}}\n{m.group(1)}\n:::',
                    t, count=1, flags=re.MULTILINE,
                )
                # Remaining *...* paragraphs (description, attribution) → Cover Text style
                t = re.sub(
                    r'^\*([^*\n]+)\*[ \t]*$',
                    lambda m: f':::{{custom-style="Cover Text"}}\n{m.group(1)}\n:::',
                    t, flags=re.MULTILINE,
                )
                # Scale cover image to fit 8.5x11 page with 1in margins
                t = re.sub(
                    r'(!\[[^\]]*\]\([^)]+\))(?!\{)',
                    r'\1{width=6.5in height=6.7in}',
                    t,
                )
                texts[i] = t

        # Write preprocessed markdown to temp files.
        # Chapter 1 gets a section break (ending the front matter section with Roman
        # page numbers) instead of a plain page break.
        tmp_files = []
        for i, (path, text) in enumerate(zip(chapter_paths, texts)):
            tmp_md = tmp / f"{i:02d}_{path.name}"
            if i == 0:
                content = text
            elif path.name == "1_IFC_Alignment_Concepts.md":
                content = _FRONT_MATTER_SECTION_BREAK + text
            else:
                content = page_break + text
            tmp_md.write_text(content, encoding="utf-8")
            tmp_files.append(str(tmp_md))

        cmd = [
            "pandoc",
            *tmp_files,
            "--from", "markdown+tex_math_dollars+raw_html-implicit_figures",
            "--to", "docx",
            f"--reference-doc={REFERENCE_DOCX}",
            f"--resource-path={ROOT}",
            f"--resource-path={tmp}",
            "--standalone",
            "--output", str(OUTPUT_DOCX),
        ]

        print(f"Running Pandoc on {len(tmp_files)} files...")
        result = subprocess.run(cmd, capture_output=True, text=True, cwd=str(ROOT))

        if result.returncode != 0:
            print("ERROR from Pandoc:")
            print(result.stderr)
            sys.exit(1)

        if result.stderr:
            svg_warns = [l for l in result.stderr.splitlines() if "svg" in l.lower()]
            other_warns = [l for l in result.stderr.splitlines() if "svg" not in l.lower()]
            if svg_warns:
                print(f"  {len(svg_warns)} remaining SVG warning(s) (images not found in source)")
            for line in other_warns:
                print(f"  {line}")

    finally:
        shutil.rmtree(tmp, ignore_errors=True)

    print("Adding page numbers...")
    add_page_numbers(OUTPUT_DOCX)

    size_mb = OUTPUT_DOCX.stat().st_size / 1_048_576
    print(f"\nOutput: {OUTPUT_DOCX}  ({size_mb:.1f} MB)")
    print()
    print("In Word, before exporting to PDF:")
    print("  1. Accept the prompt to update fields — TOC populates automatically")
    print("  2. Check math equations for OMML conversion quality")
    print("  3. File > Export > Create PDF/XPS")


if __name__ == "__main__":
    main()
